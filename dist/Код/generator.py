import json
import os
import re
import sys

from docx import Document


if getattr(sys, "frozen", False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def get_path(relative_path):
    return os.path.join(BASE_DIR, relative_path)


def load_json(path, default):
    if not os.path.exists(path):
        return default

    with open(path, "r", encoding="utf-8") as file:
        return json.load(file)


def save_json(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)

    with open(path, "w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False, indent=2)


def replace_text(doc, replacements):
    def replace_in_paragraph(paragraph):
        full_text = paragraph.text

        for key, value in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, value)

        if full_text == paragraph.text:
            return

        for run in paragraph.runs:
            run.text = ""

        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def format_car(car):
    if " - " not in car:
        return car

    name, number = car.split(" - ", 1)
    return f"{name} з номером {number}"


def sanitize_filename_part(value):
    cleaned_value = re.sub(r'[<>:"/\\|?*]', "_", value.strip())
    return cleaned_value or "ВОДІЙ"


def get_next_order_number(number_path):
    number_data = load_json(number_path, {"number": 1})
    current_number = int(number_data.get("number", 1))
    return current_number, {"number": current_number + 1}


def generate_order(driver, driver_acc, driver_dat, car, destination, date):
    settings_path = get_path("data/settings.json")
    number_path = get_path("data/order_number.json")

    settings = load_json(settings_path, {})
    template_path = get_path(settings.get("template_path", "template/order_template.docx"))

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Не знайдено шаблон документа:\n{template_path}")

    order_number, next_number_payload = get_next_order_number(number_path)
    document = Document(template_path)

    last_name = sanitize_filename_part(driver.split()[0].upper())
    mechanic_value = settings.get("mechanic", "").strip()
    signer_value = settings.get("signer", "").strip()

    replacements = {
        "{NUMBER}": str(order_number),
        "{DATE}": date,
        "{DRIVER_ACC}": driver_acc,
        "{DRIVER_DAT}": driver_dat,
        "{CAR}": format_car(car),
        "{DESTINATION}": destination,
        "{MECHANIC}": mechanic_value,
        "{SIGNER}": signer_value,
    }

    replace_text(document, replacements)

    output_folder = get_path(settings.get("output_folder", "output"))
    os.makedirs(output_folder, exist_ok=True)

    filename = f"{order_number}-г_{last_name}.docx"
    save_path = os.path.join(output_folder, filename)
    document.save(save_path)

    save_json(number_path, next_number_payload)
    return save_path
